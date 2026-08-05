import io
import markdown
from xhtml2pdf import pisa
from docx import Document
import re

def convert_markdown_to_pdf(markdown_text: str) -> io.BytesIO:
    """Convert markdown text to a PDF byte stream using xhtml2pdf."""
    html_content = markdown.markdown(markdown_text)
    
    styled_html = f"""
    <html>
    <head>
        <style>
            @page {{
                size: a4 portrait;
                margin: 2cm;
            }}
            body {{
                font-family: Helvetica, Arial, sans-serif;
                font-size: 12pt;
                line-height: 1.5;
                color: #333333;
            }}
            h1 {{ font-size: 18pt; margin-bottom: 12pt; color: #111111; }}
            h2 {{ font-size: 16pt; margin-top: 18pt; margin-bottom: 10pt; color: #222222; }}
            h3 {{ font-size: 14pt; margin-top: 14pt; margin-bottom: 8pt; }}
            p {{ margin-bottom: 10pt; }}
            ul, ol {{ margin-bottom: 10pt; }}
            li {{ margin-bottom: 4pt; }}
            hr {{ border-top: 1px solid #cccccc; margin: 20pt 0; }}
        </style>
    </head>
    <body>
        {html_content}
    </body>
    </html>
    """
    
    pdf_buffer = io.BytesIO()
    pisa.CreatePDF(io.StringIO(styled_html), dest=pdf_buffer)
    pdf_buffer.seek(0)
    
    return pdf_buffer


def convert_markdown_to_docx(markdown_text: str) -> io.BytesIO:
    """Convert markdown text to a DOCX byte stream."""
    doc = Document()
    
    lines = markdown_text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('- ') or line.startswith('* '):
            clean_text = line[2:].replace('**', '').replace('__', '').strip()
            doc.add_paragraph(clean_text, style='List Bullet')
        else:
            p = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
                    
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    
    return docx_buffer
